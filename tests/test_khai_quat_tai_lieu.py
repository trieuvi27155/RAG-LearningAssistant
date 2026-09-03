"""Test KHÁI QUÁT: các cơ chế mới phải đúng trên mọi hình dạng tài liệu, không chỉ trên
bộ tài liệu đã dùng để phát hiện ra lỗi.

Lý do có file này tách khỏi test_do_tin_cay_thuoc_do.py: hai cơ chế được thêm gần đây (đọc
lại trang PDF dính chữ, cắt bảng lớn giữ dòng tiêu đề) đều được phát hiện nhờ MỘT tài liệu
cụ thể — một cuốn sách LaTeX và một biểu mẫu Word. Nguy cơ tự nhiên là chúng được chỉnh cho
vừa đúng hai tài liệu đó rồi hỏng ở tài liệu khác, mà không ai biết.

Nên các test dưới đây cố ý dựng ra những hình dạng KHÔNG có trong bộ tài liệu test: bảng
một hàng, bảng tiêu đề dài hơn cả ngân sách, bảng toàn ô rỗng, văn bản tiếng Việt có dấu,
văn bản toàn công thức, văn bản có từ ghép dài hợp lệ. Yêu cầu chung: hoặc xử lý đúng, hoặc
suy giảm êm — tuyệt đối không mất nội dung và không tạo chunk vượt giới hạn model.
"""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import pytest

from rag.chunking import _cat_bang_giu_tieu_de, chia_chunk
from rag.document_loader import (
    MOC_BANG_DONG,
    MOC_BANG_MO,
    _bang_sang_markdown,
    _trich_text_thich_ung,
    _ty_le_dinh_chu,
    _ty_le_tu_le,
)


def _dem_tu(t: str) -> int:
    """Bộ đếm token giả lập - đếm theo từ. Dùng cho test để trần token dễ tính nhẩm."""
    return len(t.split())


# ======================================================================
# Phát hiện dính chữ: không được báo động giả trên văn bản hợp lệ
# ======================================================================

@pytest.mark.parametrize(
    "mo_ta, mau",
    [
        ("tiếng Việt có dấu", "Nhà nước là tổ chức quyền lực công cộng đặc biệt của xã hội. "),
        ("tiếng Anh thường", "The quick brown fox jumps over the lazy dog near the river. "),
        ("thuật ngữ ML dài", "Backpropagation and regularization improve generalization. "),
        (
            "văn bản kỹ thuật dày từ dài",
            "The internationalization of the interface required a "
            "counterrevolutionary change in how we handle multidimensional data. ",
        ),
        ("toàn công thức", "f(x) = 1/(1+e^-x), sigma^2 = E[(X-mu)^2], w = (X^T X)^-1 X^T y "),
        ("mã nguồn", "def tinh_diem(a, b): return sum(a) / max(len(b), 1) "),
        ("bảng markdown", "| Cột A | Cột B |\n| --- | --- |\n| giá trị 1 | giá trị 2 |\n"),
        ("toàn số", "2024 2025 2026 100 200 300 400 500 600 700 800 900 "),
    ],
)
def test_khong_bao_dong_gia_dinh_chu(mo_ta, mau):
    """Báo động giả nghĩa là trang đang đọc TỐT bị đem đi đọc lại với tham số hung hăng hơn -
    tức tự tạo ra lỗi vỡ từ ở nơi vốn không có lỗi nào.

    Nhân mẫu lên cỡ một TRANG thật (~2.000 ký tự) vì đó là thứ hàm này được gọi để đo. Một
    câu lẻ có tỉ lệ hoàn toàn khác - chính chỗ đó là nguồn báo động giả đã phát hiện được.
    """
    assert _ty_le_dinh_chu(mau * 40) < 0.10, mo_ta


@pytest.mark.parametrize(
    "mo_ta, mau",
    [
        ("tiếng Anh dính", "whichareknownasthenormalequationsfortheleastsquaresproblem "),
        ("tiếng Việt dính", "nhànướclàtổchứcquyềnlựccôngcộngđặcbiệtcủaxãhội "),
        ("dính một phần", "The result is thisisaverylongrunofgluedtogetherwordshere in text. "),
    ],
)
def test_bat_duoc_dinh_chu_o_ca_hai_ngon_ngu(mo_ta, mau):
    assert _ty_le_dinh_chu(mau * 40) >= 0.10, mo_ta


@pytest.mark.parametrize(
    "mo_ta, text",
    [
        ("chuỗi rỗng", ""),
        ("một dòng tiêu đề", "CHƯƠNG III: TRÁCH NHIỆM PHÁP LÝ"),
        ("trang chỉ có hình", "Hình 3.2"),
        ("một từ dài lẻ loi", "internationalization"),
    ],
)
def test_qua_it_chu_thi_khong_ket_luan_gi(mo_ta, text):
    """Trang gần như trống (chỉ có hình, hoặc một dòng tiêu đề) không đủ dữ liệu để tính tỉ
    lệ - và cũng chẳng có gì để mà cứu. Kết luận bừa ở đây là nguồn báo động giả lớn nhất."""
    assert _ty_le_dinh_chu(text) == 0.0, mo_ta


def test_bao_dong_gia_van_vo_hai_vi_ban_doc_lai_phai_TOT_HON_moi_duoc_nhan():
    """Tính chất an toàn quan trọng hơn mọi ngưỡng, nên phải có test riêng.

    Không ngưỡng nào đúng cho mọi tài liệu trên đời - một trang toàn từ ghép rất dài (thuật
    ngữ hoá học, tiếng Đức) vẫn có thể bị nghi oan là dính chữ. Điều khiến cơ chế này an
    toàn KHÔNG phải là ngưỡng chọn khéo, mà là: bản đọc lại chỉ được nhận khi nó ĐO ĐƯỢC là
    đỡ dính hơn VÀ không làm vỡ từ thêm. Trang không thật sự dính thì không tham số nào làm
    nó "đỡ dính" được, nên bản gốc luôn được giữ - báo động giả chỉ tốn thêm vài lần đọc,
    không bao giờ làm hỏng nội dung.

    Dựng một trang giả lập trả về CÙNG một văn bản với mọi x_tolerance (đúng hành vi của
    trang không có vấn đề về khoảng cách chữ) và kiểm tra kết quả không đổi.
    """
    van_ban = ("Deoxyribonucleicacidsequencing and electroencephalographically "
               "measured responses were analysed carefully. ") * 30

    class TrangGiaLap:
        def __init__(self):
            self.so_lan_doc = 0

        def extract_text(self, **_):
            self.so_lan_doc += 1
            return van_ban

    trang = TrangGiaLap()
    assert _ty_le_dinh_chu(van_ban) >= 0.10, "ca này phải bị nghi ngờ thì test mới có nghĩa"
    assert _trich_text_thich_ung(trang, "hoa_hoc.pdf", 1) == van_ban
    assert trang.so_lan_doc > 1, "phải có thử đọc lại thì mới chứng minh được là vô hại"


def test_ty_le_tu_le_bat_duoc_van_ban_bi_vo_tu():
    """Đây là chốt an toàn của cơ chế đọc lại - nếu phép đo này không nhạy thì tham số hung
    hăng sẽ được chấp nhận và văn bản bị phá mà không ai biết."""
    lanh_lan = "The quick brown fox jumps over the lazy dog"
    vo_vun = "T h e q u i c k b r o w n f o x"
    assert _ty_le_tu_le(lanh_lan) < 0.15
    assert _ty_le_tu_le(vo_vun) > 0.8


# ======================================================================
# Bảng: mọi hình dạng đều phải xử lý được, không mất nội dung
# ======================================================================

def _khoi(cac_dong):
    return "\n".join([MOC_BANG_MO, *cac_dong, MOC_BANG_DONG])


@pytest.mark.parametrize(
    "mo_ta, cac_dong",
    [
        ("bảng chỉ có tiêu đề", ["| A | B |", "| --- | --- |"]),
        ("bảng 1 hàng dữ liệu", ["| A | B |", "| --- | --- |", "| 1 | 2 |"]),
        ("bảng nhiều cột", ["| " + " | ".join("ABCDEFGH") + " |",
                            "| " + " | ".join(["---"] * 8) + " |",
                            "| " + " | ".join("12345678") + " |"]),
        ("bảng ô rỗng", ["| A | B |", "| --- | --- |", "| | |", "| x | |"]),
        ("bảng không có dòng gạch ngăn", ["| A | B |", "| 1 | 2 |"]),
    ],
)
def test_cat_bang_khong_lam_mat_noi_dung(mo_ta, cac_dong):
    """Bất kể hình dạng nào: mọi ô có chữ trong bảng gốc phải còn tìm thấy được sau khi cắt.
    Mất nội dung ở đây là mất âm thầm - không có lỗi nào báo ra, chỉ là câu hỏi về ô đó
    vĩnh viễn không tra được nữa."""
    khoi = _khoi(cac_dong)
    cac_manh = _cat_bang_giu_tieu_de(khoi, dem=_dem_tu, tran=12)
    gop = "\n".join(m for m, _ in cac_manh)

    for dong in cac_dong:
        for o in (x.strip() for x in dong.strip().strip("|").split("|")):
            if o and o != "---":
                assert o in gop, f"{mo_ta}: mất ô {o!r}"


def test_tieu_de_dai_hon_ngan_sach_khong_gay_lap_vo_han():
    """Nếu dòng tiêu đề tự nó đã hết ngân sách, lặp lại nó ở từng mảnh sẽ đẩy MỌI hàng dữ
    liệu ra ngoài. Phải nhận ra và lùi về cách cắt cũ, thay vì sinh ra một loạt mảnh chỉ
    chứa tiêu đề."""
    tieu_de = "| " + " | ".join(f"Tên cột rất dài số {i}" for i in range(10)) + " |"
    khoi = _khoi([tieu_de, "| " + " | ".join(["---"] * 10) + " |",
                  "| " + " | ".join(str(i) for i in range(10)) + " |"])

    cac_manh = _cat_bang_giu_tieu_de(khoi, dem=_dem_tu, tran=15)
    assert len(cac_manh) == 1, "không lặp được tiêu đề thì phải trả nguyên khối, không cắt vụn"


def test_khong_chunk_nao_vuot_gioi_han_model_du_bang_hinh_dang_nao():
    """Ràng buộc CỨNG của cả hệ thống: chunk vượt max_seq_length bị model cắt ÂM THẦM lúc
    encode, nên phần cuối không bao giờ tra được. Bảng là chỗ dễ vi phạm nhất vì nó được ưu
    tiên giữ nguyên khối."""
    tieu_de_dai = "| " + " | ".join(f"Cột có tên dài dòng số {i}" for i in range(12)) + " |"
    cac_trang = [
        {"nguon": "a.docx", "trang": 1, "noidung": _khoi(
            [tieu_de_dai, "| " + " | ".join(["---"] * 12) + " |"]
            + ["| " + " | ".join(f"giá trị {i}-{j}" for j in range(12)) + " |" for i in range(30)]
        )},
        {"nguon": "b.docx", "trang": 1, "noidung": _khoi(
            ["| Mục | Nội dung |", "| --- | --- |", "| 1 | " + "văn xuôi rất dài " * 200 + "|"]
        )},
    ]
    cac_chunk = chia_chunk(cac_trang, dem_token_fn=_dem_tu, max_seq_length=None)

    assert cac_chunk
    for c in cac_chunk:
        assert _dem_tu(c["noidung"]) <= 160, f"chunk {_dem_tu(c['noidung'])} token vượt trần"


def test_bang_vua_gioi_han_van_duoc_giu_nguyen_khoi():
    """Không được cắt bảng nhỏ chỉ vì có cơ chế cắt: bảng vừa giới hạn phải đi nguyên khối,
    đúng như quyết định ở §5.32."""
    khoi = _khoi(["| A | B |", "| --- | --- |", "| 1 | 2 |", "| 3 | 4 |"])
    cac_chunk = chia_chunk(
        [{"nguon": "a.pptx", "trang": 1, "noidung": khoi}], dem_token_fn=_dem_tu
    )
    cac_bang = [c for c in cac_chunk if c["loai_noi_dung"] == "bang"]
    assert len(cac_bang) == 1
    assert "| 1 | 2 |" in cac_bang[0]["noidung"] and "| 3 | 4 |" in cac_bang[0]["noidung"]


# ======================================================================
# Bỏ cột rỗng: phải giữ nguyên dữ liệu, chỉ bỏ cột KHÔNG có gì
# ======================================================================

def test_chi_bo_cot_rong_o_moi_hang():
    bang = [["A", "", "B", ""], ["1", "", "2", ""], ["3", "", "4", ""]]
    md = _bang_sang_markdown(bang)
    assert md.splitlines()[0] == "| A | B |"
    for gia_tri in ("1", "2", "3", "4"):
        assert f"| {gia_tri} " in md or f" {gia_tri} |" in md


def test_khong_bo_cot_chi_rong_o_MOT_SO_hang():
    """Ô trống rải rác là dữ liệu thật (biểu mẫu bỏ trống mục), không phải cột thừa."""
    bang = [["A", "B"], ["1", ""], ["", "2"]]
    md = _bang_sang_markdown(bang)
    assert md.splitlines()[0] == "| A | B |", "bảng 2 cột phải giữ đủ 2 cột"


def test_bang_toan_o_rong_khong_gay_loi():
    assert _bang_sang_markdown([["", ""], ["", ""]]) == ""
    assert _bang_sang_markdown([]) == ""


# ======================================================================
# Bố cục NHIỀU CỘT — đọc ngang trang 2 cột sẽ trộn câu của hai cột
# ======================================================================

def _tao_pdf_hai_cot(duong_dan):
    """Dựng một PDF 2 cột thật bằng reportlab (đã là dependency của evaluation)."""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    trai = ["Dieu 1. Pham vi dieu chinh cua luat",
            "bao gom cac quan he phat sinh trong",
            "hoat dong dau tu kinh doanh tren lanh",
            "tho quoc gia theo quy dinh hien hanh"]
    phai = ["Dieu 2. Doi tuong ap dung gom moi",
            "doanh nghiep duoc thanh lap va hoat",
            "dong theo quy dinh cua Luat Doanh",
            "nghiep khong phan biet thanh phan"]

    c = canvas.Canvas(str(duong_dan), pagesize=A4)
    rong, cao = A4
    c.setFont("Helvetica", 10)
    for k in range(5):
        for i, dong in enumerate(trai):
            c.drawString(60, cao - 90 - (k * 4 + i) * 16, dong)
        for i, dong in enumerate(phai):
            c.drawString(rong / 2 + 25, cao - 90 - (k * 4 + i) * 16, dong)
    c.save()


def test_trang_hai_cot_khong_bi_tron_cau_cua_hai_cot(tmp_path, monkeypatch):
    """Đây là hỏng ngay từ khâu đọc: pdfplumber đọc theo dòng ngang suốt bề ngang trang, nên
    câu của cột trái bị nối thẳng vào câu của cột phải. Mọi chunk sinh ra đều vô nghĩa, và
    không có dấu hiệu nào để nhận ra ngoài việc câu trả lời lộn xộn."""
    import config
    from rag.document_loader import doc_pdf

    monkeypatch.setattr(config, "BAT_TRICH_ANH", False)
    monkeypatch.setattr(config, "BAT_OCR_DU_PHONG", False)

    f = tmp_path / "hai_cot.pdf"
    _tao_pdf_hai_cot(f)
    noidung = doc_pdf(f)[0]["noidung"]

    assert "cua luat Dieu 2." not in noidung, "câu của hai cột vẫn bị dính vào nhau"
    assert "Dieu 1. Pham vi dieu chinh cua luat" in noidung
    assert "Dieu 2. Doi tuong ap dung gom moi" in noidung


def test_trang_mot_cot_khong_bi_nhan_nham_thanh_nhieu_cot(tmp_path, monkeypatch):
    """Chốt an toàn quan trọng hơn: nhận nhầm trang 1 cột thành 2 cột sẽ CẮT ĐÔI một tài liệu
    vốn đang đọc tốt. Bản dò đầu tiên thiếu chốt này đã nhận nhầm 100% số trang của một giáo
    trình thật, vì nó tưởng LỀ TRANG là rãnh giữa cột."""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    import config
    from rag.document_loader import _cac_cot_cua_trang

    import pdfplumber

    monkeypatch.setattr(config, "BAT_TRICH_ANH", False)
    f = tmp_path / "mot_cot.pdf"
    c = canvas.Canvas(str(f), pagesize=A4)
    rong, cao = A4
    c.setFont("Helvetica", 11)
    for i in range(30):
        c.drawString(70, cao - 80 - i * 18,
                     "Day la mot dong van xuoi binh thuong cua trang mot cot trong tai lieu")
    c.save()

    with pdfplumber.open(f) as pdf:
        assert _cac_cot_cua_trang(pdf.pages[0]) == [], "trang 1 cột bị nhận nhầm thành nhiều cột"


# ======================================================================
# TEXT BOX trong DOCX — python-docx không thấy, nội dung mất hẳn
# ======================================================================

def test_noi_dung_trong_text_box_docx_khong_bi_mat(tmp_path, monkeypatch):
    """Trong Word, sơ đồ / khung "Lưu ý" / trích dẫn nổi bật hay nằm trong text box - và đó
    thường là chỗ cô đọng nhất của trang. `Paragraph.text` của python-docx trả về chuỗi rỗng
    cho chúng, nên nội dung biến mất khỏi index mà không có dấu hiệu gì."""
    from docx import Document
    from docx.oxml import parse_xml

    import config
    from rag.document_loader import doc_docx

    monkeypatch.setattr(config, "BAT_TRICH_ANH", False)

    d = Document()
    d.add_paragraph("Doan van thuong nam ngoai text box.")
    run = d.add_paragraph().add_run()
    run._element.append(parse_xml(
        '<w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        "<w:txbxContent><w:p><w:r><w:t>Quy trinh gom ba buoc chinh</w:t></w:r></w:p>"
        "</w:txbxContent></w:drawing>"
    ))
    f = tmp_path / "co_text_box.docx"
    d.save(f)

    # python-docx tự nó KHÔNG thấy nội dung này - đó là lý do phải tự dò XML.
    assert all("ba buoc" not in p.text for p in Document(f).paragraphs)

    noidung = doc_docx(f)[0]["noidung"]
    assert "Quy trinh gom ba buoc chinh" in noidung
    assert "Doan van thuong nam ngoai text box." in noidung
